# Import the necessary libraries
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches
from docx.shared import Pt, RGBColor
from docx.oxml.xmlchemy import OxmlElement
from docx.oxml.shared import qn

from src.interface import Project, Interface, ParseAPIDoc


# Define a class for the API documentation
class APIDocumentation:

    def _set_cell_background(self, cell, fill, color=None, val=None):
        """
        @fill: Specifies the color to be used for the background
        @color: Specifies the color to be used for any foreground
        pattern specified with the val attribute
        @val: Specifies the pattern to be used to lay the pattern
        color over the background color.
        """
        cell_properties = cell._element.tcPr
        try:
            cell_shading = cell_properties.xpath('w:shd')[0]  # in case there's already shading
        except IndexError:
            cell_shading = OxmlElement('w:shd')  # add new w:shd element to it
        if fill:
            cell_shading.set(qn('w:fill'), fill)  # set fill property, respecting namespace
        if color:
            pass  # TODO
        if val:
            pass  # TODO
        cell_properties.append(cell_shading)  # finally extend cell props with shading element

    def __init__(self, project: Project, title=None):
        self.project = project
        self.doc_name = title if title else project.title
        # Create a new document
        self.document = Document()
        self.document.add_heading(self.doc_name, level=0)

        for group in project.interface_group_list:
            self.create_group(group)
            for interface in group.interface_list:
                self.create_interface(interface)

    def create_group(self, group):
        self.document.add_heading(group.title, 1)


    def create_request_param(self, request_params):
        # Add the subheadings for the request and response parameters
        self.document.add_heading('请求参数', level=3)
        table = self.document.add_table(rows=1, cols=6)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER  # 表格居中
        table.autofit = False
        col_width_dic = {
            0: 1,
            1: 3,
            2: 1,
            3: 1,
            4: 1,
            5: 3
        }
        for col_num in range(6):
            table.cell(0, col_num).width = Inches(col_width_dic[col_num])

        table_style = self.document.styles['Table Grid']
        table.style = table_style
        # table.style.cell_background_color = WD_COLOR_INDEX
        table.style.font.size = 10

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '参数名'
        hdr_cells[1].text = '中文说明'
        hdr_cells[2].text = '数据类型'
        hdr_cells[3].text = '参数类型'
        hdr_cells[4].text = '是否必填'
        hdr_cells[5].text = '备注'
        for cell in hdr_cells:
            self._set_cell_background(cell, "BDD6EE")
            cell.paragraphs[0].style.font.size = Pt(10)

        for param in request_params:

            row_cells = table.add_row().cells

            row_cells[0].text = param.deep * self.project.separator + param.name
            row_cells[1].text = param.desc
            row_cells[2].text = "object" if param.schema_value else param.type
            row_cells[3].text = param.mode if param.mode else ""
            row_cells[4].text = '是' if param.require else '否'
            row_cells[5].text = ""

    def create_response_param(self, response_params):

        self.document.add_heading('响应参数', level=3)
        table = self.document.add_table(rows=1, cols=3)
        table_style = self.document.styles['Table Grid']
        table.style = table_style
        # table.style.cell_background_color = WD_COLOR_INDEX
        table.style.font.size = 10

        table.alignment = WD_TABLE_ALIGNMENT.CENTER  # 表格居中
        table.autofit = False
        col_width_dic = {
            0: 2,
            1: 2,
            2: 6,
        }
        for col_num in range(3):
            table.cell(0, col_num).width = Inches(col_width_dic[col_num])

        hdr_cells = table.rows[0].cells

        hdr_cells[0].text = '参数名'
        hdr_cells[1].text = '参数类型'
        hdr_cells[2].text = '说明'
        for cell in hdr_cells:
            self._set_cell_background(cell, "BDD6EE")
            cell.paragraphs[0].style.font.size = Pt(10)

        for param in response_params:
            resp_map = project.resp_map.get(param.name, {})
            project.resp_map.get('desc', param.desc)
            row_cells = table.add_row().cells
            row_cells[0].text = param.deep * self.project.separator + param.name
            row_cells[1].text = "object" if param.schema_value else param.type
            row_cells[2].text = resp_map.get("chinese", param.desc) + resp_map.get('desc', "")

    def create_basic_information(self, interface: Interface):
        # Add the main title
        self.document.add_heading(interface.name, 2)

        # Add the subheadings
        self.document.add_paragraph(interface.desc)

        self.document.add_heading('基本信息', level=3)

        # Add the details for the basic information
        self.document.add_paragraph('接口地址：' + interface.url)
        self.document.add_paragraph('请求方式：' + interface.method)


    def create_interface(self, interface: Interface):
        self.create_basic_information(interface)
        self.create_request_param(interface.request_params)
        self.create_response_param(interface.response_params)

    def set_font(self):
        # 获取所有段落和文本框
        elements = self.document.paragraphs

        # 遍历所有元素
        for element in elements:
            # 获取元素的字体
            font = element.style.font
            # 将字体设置为宋体
            font.name = '宋体'

        # 遍历所有标题
        for heading in self.document.paragraphs:
            # 获取标题的字体
            font = heading.style.font
            # 将字体设置为宋体
            font.name = '宋体'

        # 遍历所有表格
        for table in self.document.tables:
            # 遍历表格的所有单元格
            for row in table.rows:
                for cell in row.cells:
                    # 获取单元格的字体
                    font = cell.paragraphs[0].style.font
                    # 将字体设置为宋体
                    font.name = '宋体'

    def save_document(self):
        # Save the document
        self.set_font()
        self.document.save(self.doc_name + '.docx')


if __name__ == '__main__':
    project = ParseAPIDoc().run('swagger.html')

    api_doc = APIDocumentation(project, "分析研判")
    api_doc.save_document()
